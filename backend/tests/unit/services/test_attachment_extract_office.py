"""Office-document extraction tests for ``services.knowledge``.

Covers the DOCX + XLSX paths added alongside the chat-attachment ingest
fix. We synthesize tiny in-memory documents so the test doesn't depend
on shipping fixture binaries with the repo.
"""

from __future__ import annotations

import io
import uuid

import pytest

from app.db.models.attachment import Attachment, AttachmentKind
from app.services.knowledge import (
    AttachmentExtractError,
    extract_text_from_attachment,
)


def _make_att(
    *, filename: str, mime_type: str, size: int, kind: AttachmentKind = AttachmentKind.DOCUMENT
) -> Attachment:
    """Build a transient ``Attachment`` row sufficient for the extractor."""
    a = Attachment(
        id=uuid.uuid4(),
        workspace_id=uuid.uuid4(),
        filename=filename,
        mime_type=mime_type,
        size_bytes=size,
        kind=kind,
        storage_uri="/tmp/missing.bin",
    )
    return a


class TestDocx:
    def _build_docx_bytes(self) -> bytes:
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        doc = Document()
        doc.add_heading("登革热媒介伊蚊孳生地清理指引", level=1)
        doc.add_paragraph("本指引覆盖白纹伊蚊与埃及伊蚊的孳生环境识别与清理。")
        doc.add_paragraph("")  # blank paragraph (should be skipped)
        doc.add_paragraph("第一节：常见孳生地。")  # noqa: RUF001 - intentional Chinese punctuation
        # Add a table with two rows.
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "类型"
        table.rows[0].cells[1].text = "清理方法"
        table.rows[1].cells[0].text = "废弃轮胎"
        table.rows[1].cells[1].text = "翻转、打孔或回收"
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_extracts_paragraphs_and_tables(self):
        data = self._build_docx_bytes()
        att = _make_att(
            filename="guide.docx",
            mime_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            size=len(data),
        )
        out = extract_text_from_attachment(att, data)
        assert "登革热媒介伊蚊孳生地清理指引" in out
        assert "白纹伊蚊与埃及伊蚊" in out
        assert "类型 | 清理方法" in out
        assert "废弃轮胎 | 翻转、打孔或回收" in out

    def test_legacy_doc_binary_raises_clear_error(self):
        """Old binary .doc isn't supported — surface a clear error code."""
        att = _make_att(
            filename="legacy.doc",
            mime_type="application/msword",
            size=4,
        )
        with pytest.raises(AttachmentExtractError) as exc:
            extract_text_from_attachment(att, b"\xd0\xcf\x11\xe0")  # OLE2 magic
        assert exc.value.code == "docx_parse_failed"

    def test_image_only_docx_raises_empty(self):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        doc = Document()
        # No paragraphs / tables — empty body.
        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
        att = _make_att(
            filename="empty.docx",
            mime_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            size=len(data),
        )
        with pytest.raises(AttachmentExtractError) as exc:
            extract_text_from_attachment(att, data)
        assert exc.value.code == "docx_empty"


class TestXlsx:
    def _build_xlsx_bytes(self) -> bytes:
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")
        wb = Workbook()
        ws = wb.active
        ws.title = "孳生地"
        ws.append(["类型", "示例", "数量"])
        ws.append(["容器积水", "花盆托盘", 12])
        ws.append([None, None, None])  # blank row → skipped
        ws.append(["废弃物", "塑料瓶", 4])
        ws2 = wb.create_sheet("清理方法")
        ws2.append(["措施", "频次"])
        ws2.append(["翻转倒扣", "每周"])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_extracts_each_sheet(self):
        data = self._build_xlsx_bytes()
        att = _make_att(
            filename="data.xlsx",
            mime_type=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
            size=len(data),
        )
        out = extract_text_from_attachment(att, data)
        assert "# Sheet: 孳生地" in out
        assert "类型\t示例\t数量" in out
        assert "容器积水\t花盆托盘\t12" in out
        assert "# Sheet: 清理方法" in out
        assert "翻转倒扣\t每周" in out
